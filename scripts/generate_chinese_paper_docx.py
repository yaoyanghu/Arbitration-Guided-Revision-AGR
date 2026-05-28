from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO_ROOT = Path(r"D:\HUYAOYANG\Work\New_ChronoRAG")
OUTPUT_DIR = Path(r"D:\HUYAOYANG\Work\new")
REPO_DOCX = REPO_ROOT / "reports" / "final_refresh" / "NEW_CHRONORAG_CHINESE_PAPER_DRAFT_20260328.docx"
OUTPUT_DOCX = OUTPUT_DIR / "New_ChronoRAG_中文论文草稿_20260328.docx"


def set_fonts(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_para(document: Document, text: str, *, bold: bool = False) -> None:
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.28)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    run.bold = bold


def add_table(document: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    if title:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.add_paragraph()


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    REPO_DOCX.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_fonts(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    add_title(doc, "面向时序与冲突证据的可信答案级检索增强生成")

    doc.add_heading("摘要", level=1)
    add_para(doc, "检索增强生成系统在存在过时证据、当前证据与冲突证据的场景下，往往难以保证答案选择的时效一致性和证据一致性。本文围绕一个收口后的 New_ChronoRAG 主线，研究答案级 faithful RAG 在时序约束与冲突证据条件下的可解释改进。我们将最终系统冻结为检索、temporal-aware answer extraction 和 conflict-aware evidence arbitration 三部分，并采用分层数据设计：HOH 作为主结果集，TempRAGEval 作为时间敏感辅助验证集，FEVER official retrieval 作为受控辅助检索基准。结果表明，在 HOH formal 1024 上，完整模型相较于 stronger retrieval baseline、no_temporal 和 no_conflict 均取得更优的答案级表现；在 TempRAGEval formal 1244 上，完整模型相较于 no_temporal 保持稳定优势，并显著提升 citation title recall。进一步地，我们系统审计了结构化判别层、轻量 learned scorer 和可信生成/拒答三类后续增强。结果显示，结构化层和 learned scorer 未带来独立主线增益，而可信生成虽然降低了 EM/F1，却能有效抑制 unsupported answers，适合作为审稿防御型扩展。本文最终主张保持克制：已被真实验证的核心贡献仅包括 temporal-aware answer extraction 与 conflict-aware evidence arbitration，而 source/reliability 不作为本文已证实的核心贡献。")

    doc.add_heading("关键词", level=1)
    add_para(doc, "检索增强生成；时序问答；冲突证据；可信生成；答案级评测；证据仲裁")

    doc.add_heading("1 引言", level=1)
    add_para(doc, "在真实知识环境中，问题往往对应的不只是相关或不相关证据，还包括时间上已经过时但词面高度相似的旧证据、与当前答案值冲突的同标题证据，以及同一证据家族内部高度相似但实体不同的统计型句子。此时，仅依靠检索分数或浅层 lexical fit 容易产生表面合理、实则不被当前证据支持的答案。")
    add_para(doc, "New_ChronoRAG 的目标不是提出一个庞杂的新架构，而是在可复现、可解释、可审计的前提下，构建一个面向时序与冲突证据的答案级 faithful RAG 管线。本文将最终主线收口为检索、temporal-aware answer extraction 与 conflict-aware evidence arbitration 三部分，以保证 claim 与 evidence 严格对齐。")

    doc.add_heading("2 相关工作", level=1)
    add_para(doc, "相关工作主要包括检索增强生成、时间敏感问答、冲突证据处理以及可信输出控制。本文与这些工作的区别不在于追求更大规模的生成模型，而在于对答案级 faithful RAG 问题进行收口式建模与验证，并对不成功的后续增强保留诚实报告。")

    doc.add_heading("3 方法", level=1)
    add_para(doc, "本文最终方法由三个阶段组成：检索、temporal-aware answer extraction，以及 conflict-aware evidence arbitration。系统首先检索少量候选证据，再根据问题时间关系和答案类型选择更合适的答案承载句，最后在 stale/current 与冲突证据之间进行仲裁。")
    add_para(doc, "图 1 方法流程图建议：问题 -> 检索候选证据 -> temporal-aware answer extraction -> conflict-aware evidence arbitration -> citation-aware answer -> （可选）trustworthy abstention。")
    add_para(doc, "在后续增强中，我们进一步测试了结构化判别层、轻量 learned scorer 以及可信生成/拒答扩展，但仅将真实有效的部分保留到最终叙事中。")

    doc.add_heading("4 实验设置", level=1)
    add_para(doc, "本文使用分层数据设计：HOH 为主结果集，TempRAGEval 为时间敏感辅助验证集，FEVER official retrieval 为受控辅助检索基准，Route A 保留为诊断资产，Route B 冻结不回主线。Formal 主表方法共享同一评测切片和预算。")
    add_para(doc, "主要评测指标包括 EM、token F1 和 citation title recall。在可信生成扩展中，额外记录 abstention rate 以及 refusal / abstention audit。")

    doc.add_heading("5 主结果", level=1)
    add_para(doc, "表 1 展示 HOH formal 主结果。完整模型在 EM 和 token F1 上均优于 stronger retrieval baseline、no_temporal 和 no_conflict，表明 temporal-aware answer extraction 与 conflict-aware evidence arbitration 均对最终答案级系统有独立帮助。")
    add_table(
        doc,
        "表 1 HOH Formal Main Table",
        ["方法", "EM", "token F1", "citation title recall"],
        [
            ["stronger_retrieval_template", "0.184", "0.273", "0.947"],
            ["no_temporal", "0.192", "0.286", "0.947"],
            ["no_conflict", "0.200", "0.303", "0.945"],
            ["full_model", "0.207", "0.310", "0.945"],
        ],
    )
    add_para(doc, "表 2 展示 TempRAGEval formal 辅助结果。完整模型在 EM、token F1 和 citation title recall 上均优于 no_temporal，说明 temporal-aware extraction 的收益并非只存在于主结果集内部。")
    add_table(
        doc,
        "表 2 TempRAGEval Formal Auxiliary Table",
        ["方法", "EM", "token F1", "citation title recall"],
        [
            ["vanilla_rag_extractive", "0.000", "0.127", "0.344"],
            ["stronger_retrieval_template", "0.064", "0.104", "0.342"],
            ["no_temporal", "0.051", "0.095", "0.344"],
            ["full_model", "0.076", "0.122", "0.400"],
        ],
    )
    add_para(doc, "表 3 给出 FEVER controlled auxiliary retrieval 结果，用于保留检索层 continuity，而不是作为答案级主张的核心证据。")
    add_table(
        doc,
        "表 3 FEVER Controlled Auxiliary Table",
        ["setting", "strict@1", "strict@5", "strict@10", "relaxed@1", "relaxed@5", "relaxed@10"],
        [
            ["baseline BM25", "0.368", "0.627", "0.720", "0.737", "0.883", "0.918"],
            ["BM25 + title overlap rerank", "0.415", "0.680", "0.720", "0.760", "0.895", "0.918"],
        ],
    )

    doc.add_heading("6 外部竞品与消融", level=1)
    add_para(doc, "我们在 HOH 和 TempRAGEval 上进一步加入外部风格竞品，包括 HyDE-like 与 CRAG-like 的轻量、公平实现。结果表明，这些强检索对照能逼近 stronger retrieval baseline，但仍未追平完整模型。")
    add_table(
        doc,
        "表 4 External Competitor Table (HOH)",
        ["方法", "EM", "token F1", "citation title recall"],
        [
            ["vanilla_rag_extractive", "0.000", "0.203", "0.946"],
            ["stronger_retrieval_template", "0.181", "0.271", "0.946"],
            ["hyde_like", "0.182", "0.271", "0.947"],
            ["crag_like", "0.182", "0.270", "0.943"],
            ["full_model", "0.202", "0.306", "0.946"],
        ],
    )
    add_para(doc, "最小正式消融显示：temporal 在 HOH 和 TempRAGEval 上都被验证，conflict 在 HOH 上被验证，而 source/reliability 没有被保留为核心贡献。")

    doc.add_heading("7 Case Study", level=1)
    add_para(doc, "Temporal success cases 主要表现为系统能够忽略显眼但不满足 query-time relation 的年份，而返回真正与 as of、last time、most recent 等关系一致的年份。Conflict success cases 主要体现在同标题或同证据家族的 stale/current 冲突条件下，系统更稳定地偏向 fresher consistent evidence。")
    add_table(
        doc,
        "表 5 Case Study Summary Table",
        ["类型", "数量", "说明"],
        [
            ["temporal success", "3", "正确处理回顾性年份与 query-time relation"],
            ["conflict success", "3", "正确处理 stale/current 与 same-title conflicts"],
            ["representative failure", "3", "主要集中在 nearby-entity 与 same-format confusion"],
        ],
    )

    doc.add_heading("8 Error Analysis", level=1)
    add_para(doc, "HOH 上的主要错误已经不再是完全找不到证据，而是 same-format statistic sentence confusion、nearby-entity confusion、slot/role confusion，以及 stale/current under similar lexical fit。TempRAGEval 上的主要问题则从显眼年份误提取，逐渐转向 retrieval precision 和 answer-bearing sentence 选择不稳。")
    add_table(
        doc,
        "表 6 Error Taxonomy Summary Table",
        ["错误类型", "主要数据集", "说明"],
        [
            ["same-format statistic confusion", "HOH", "相邻实体的统计句格式相同，值选择不稳"],
            ["nearby-entity confusion", "HOH", "问题实体锚点不足时证据家族选错"],
            ["relation-sensitive temporal mistake", "TempRAGEval", "年份合理但不满足问题关系"],
            ["evidence sufficiency failure", "HOH/TempRAGEval", "第一证据已足够时第二证据仍注入噪声"],
        ],
    )

    doc.add_heading("9 Follow-up Negative Results and Trustworthiness Extension", level=1)
    add_para(doc, "在 formal mainline 完成后，我们继续测试了三类增强：结构化判别层、轻量 learned scorer 和可信生成/拒答。结构化层在 HOH 256 上相对于 no_structured 没有独立增益；轻量 learned scorer 在 HOH 256 上相对于 no_learned 反而略有下降；因此二者不进入主线。")
    add_para(doc, "可信生成扩展通过句级支持检查、unsupported span suppression 和低置信度 abstention 来提升 faithfulness。它降低了 EM/F1，但显著减少了 unsupported outputs，更适合作为 reviewer-defense 增强。")
    add_table(
        doc,
        "表 7 Trustworthy Generation Audit Table",
        ["指标", "数值"],
        [
            ["abstention rate", "0.105"],
            ["abstentions total", "27"],
            ["abstentions blocking previously wrong answers", "22"],
            ["abstentions suppressing previously correct answers", "5"],
        ],
    )

    doc.add_heading("10 Discussion", level=1)
    add_para(doc, "本文的主要优势不在于提出一个巨大复杂的新模型，而在于给出一条收口清晰、证据链完整、可审计的 faithful answer-level RAG 主线。我们对有效模块和无效模块都做了真实记录，这提高了整体叙事的可信度。")

    doc.add_heading("11 Limitations", level=1)
    add_para(doc, "本文系统仍然是中等复杂度、启发式较强的 faithful RAG 管线，而不是高度学习化的大模型推理系统。conflict 的独立收益主要由 HOH 主结果集支撑；source/reliability 虽已接通但没有显示独立 answer-level gain；结构化层和 learned scorer 没有改善主线结果；可信生成扩展虽然减少 unsupported answers，但会带来 EM/F1 回落。")

    doc.add_heading("12 Conclusion", level=1)
    add_para(doc, "本文围绕收口后的 New_ChronoRAG 主线，系统研究了时序与冲突证据条件下的答案级 faithful RAG。正式实验表明，temporal-aware answer extraction 与 conflict-aware evidence arbitration 分别在主结果集与辅助集上展现了真实价值。与此同时，我们对 source、结构化判别层、learned scorer 和可信生成扩展进行了系统 follow-up，并通过 stop-loss 保持论文叙事克制。")

    doc.add_heading("参考文献", level=1)
    references = [
        "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS, 2020.",
        "Izacard G, Grave E. Leveraging Passage Retrieval with Generative Models for Open Domain Question Answering. ACL, 2021.",
        "Gao L, Ma X, Lin J, Callan J. Precise Zero-Shot Dense Retrieval without Relevance Labels. ACL, 2023.",
        "Asai A, Wu Z, Wang Y, et al. Self-RAG: Learning to Retrieve, Generate, and Critique through Self-Reflection. ICLR, 2024.",
        "Yan S, Yu H, Guo Z, et al. Corrective Retrieval Augmented Generation. arXiv, 2024.",
        "Thorne J, Vlachos A, Christodoulopoulos C, Mittal A. FEVER: a Large-scale Dataset for Fact Extraction and Verification. NAACL, 2018.",
        "Yang Z, Qi P, Zhang S, et al. HotpotQA: A Dataset for Diverse, Explainable Multi-hop Question Answering. EMNLP, 2018.",
    ]
    for ref in references:
        add_para(doc, ref)

    doc.save(REPO_DOCX)
    doc.save(OUTPUT_DOCX)
    print(f"saved_repo={REPO_DOCX}")
    print(f"saved_output={OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
